import numpy as np
import csv
from models import *
import matplotlib.pyplot as plt
from docx import Document
from collections import Counter
import time

start_time = time.time()

labels,bitarrays = read_csv('dataset2.csv')

x1 = np.array(bitarrays)
y1 = np.array(labels)

# Mapping (for verification or later use)
# Get unique labels (your classes)

# Convert labels in y1 to integer indices

unique_labels = mapping(y1)
label_to_index = {label: idx for idx, label,in enumerate(unique_labels)}
index_to_label = {idx: label for label,idx in label_to_index.items()}

y_indices= np.array([label_to_index[label]for label in y1])

num_classes = len(unique_labels)
y_one_hot = np.eye(num_classes)[y_indices]  # shape: (num_samples, 62)

# Activation
def sigmoid(x):
    return 1 / (1 + np.exp(-x))

# Weight generator
def generate_wt(x, y):
    return np.random.randn(x, y)

# Loss function
def loss(out, Y):
    return np.sum(np.square(out - Y)) / len(Y)

def add_bitmap_to_label_unique(filename, label, bitmap):
    bitmap_str = ",".join(str(bit) for bit in bitmap)
    entry_exists = False

    # Step 1: Check for duplicates
    with open(filename, mode='r', newline='') as file:
        reader = csv.reader(file)
        for row in reader:
            if len(row) < 2:
                continue
            if row[0] == label and row[1].strip() == bitmap_str:
                entry_exists = True
                break

    if not entry_exists:
        with open(filename, mode='a', newline='') as file:
            writer = csv.writer(file)
            writer.writerow([label, bitmap_str])

def fnn_forward(input,weights1,weights2):
    h1_in = np.dot(input,weights1)
    h1_out = sigmoid(h1_in)

    h2_in = np.dot(h1_out,weights2)
    h2_out = sigmoid(h2_in)

    return h2_out,h1_out

def fnn_backprop(data,labels,weights1,weights2,lr):
    h2_out,h1_out = fnn_forward(data,weights1,weights2)

    output_loss = h2_out - labels
    hidden_loss = np.multiply(np.dot(output_loss,weights2.T), h1_out * (1 - h1_out))

    weights1_adj = np.outer(data,hidden_loss)
    weights2_adj = np.outer(h1_out,output_loss)

    weights1 -= lr * weights1_adj
    weights2 -= lr * weights2_adj

    return weights1,weights2

def train_fnn(input,labels,weights1,weights2,lr,epochs):
    accuracies = []
    losses = []
    for epoch in range(epochs):
        l = []
        correct = 0
        for i in range(len(input)):
            out,_ = fnn_forward(input[i],weights1,weights2)
            loss = mean_square(out,labels[i])
            l.append(loss)
            if np.argmax(out) == np.argmax(labels[i]):
                correct += 1
            weights1,weights2 = fnn_backprop(input[i],labels[i],weights1,weights2,lr)

        avg_loss = sum(l)/len(input)
        accuracy = (correct/len(input)) * 100
        print(f"Epoch {epoch + 1}: Accuracy = {accuracy:.2f}% , Avg Loss = {avg_loss:.4f}")
        
        accuracies.append(accuracy)
        losses.append(avg_loss)
    return accuracies,losses

def fnn_predict(input,weights1,weights2,label,axis = 1):
    if axis == 1:
        input = input[0]
    
    original_label = label
    out,_ = fnn_forward(input,weights1,weights2)
    out_percentage = out * 100
    max_value = np.max(out_percentage)
    label_idx = np.argmax(max_value)
    return label_idx

doc = Document()
doc.add_heading("FFNN Classification Report", 0)

HLN_values = [40] 
I_P_N = 187
# H_L = 40
O_P_N = 62
# Initialize weights
for H_L in HLN_values:
    doc.add_paragraph(f"==========Hidden Layer Neurons:{H_L}==========")
    w1 = generate_wt(I_P_N, H_L)
    w2 = generate_wt(H_L, O_P_N)

    # Train the model
    acc, losss, w1, w2 = train_fnn(x1, y_one_hot, w1, w2, 0.1, epochs=90)

    seen = set()
    for idx, label in enumerate(y1):
        if label not in seen:
            seen.add(label)
            data = x1[idx]
            predicted_idx = fnn_predict(data,w1,w2,label)
            predicted_label = label_to_index(int(predicted_idx))
    

    end_time = time.time()

    doc.add_paragraph(f"Time taken to classify all the characters with noisy vectors through FFNN is: {end_time - start_time:.2f}seconds")

a2,a1 = fnn_forward(x1, w1, w2)

# doc.save(f"output_with_updation_fnn.docx")
print(f"Results saved to 'output.docx'")
 

# Plot Accuracy
plt.plot(acc)
plt.ylabel("Accuracy")
plt.xlabel("Epochs")
plt.title("Training Accuracy")
plt.show()

# Plot Loss
plt.plot(losss)
plt.ylabel("Loss")
plt.xlabel("Epochs")
plt.title("Training Loss")
plt.show()